"""Shared Word doc scaffolding for event-study findings reports.

Every brief produces a 7-section findings doc. This module provides the
skeleton; the brief-specific script supplies the prose / numbers.

Public API:
    open_doc(title) -> Document
    add_section(doc, num, title, paragraphs=None, bullets=None)
    add_table(doc, headers, rows, style="Light Grid Accent 1")
    add_picture(doc, path, width_inches=6.0)
    save_doc(doc, path)
"""
from __future__ import annotations

from pathlib import Path
from typing import Iterable

from docx import Document
from docx.shared import Inches


def open_doc(title: str) -> Document:
    doc = Document()
    doc.add_heading(title, level=0)
    return doc


def add_section(doc: Document, num: int, title: str,
                paragraphs: Iterable[str] | None = None,
                bullets: Iterable[str] | None = None) -> None:
    doc.add_heading(f"{num}. {title}", level=1)
    for p in (paragraphs or []):
        doc.add_paragraph(p)
    for b in (bullets or []):
        doc.add_paragraph(b, style="List Bullet")


def add_subsection(doc: Document, title: str,
                   paragraphs: Iterable[str] | None = None) -> None:
    doc.add_heading(title, level=2)
    for p in (paragraphs or []):
        doc.add_paragraph(p)


def add_table(doc: Document, headers: list[str], rows: Iterable[Iterable],
              style: str = "Light Grid Accent 1") -> None:
    tbl = doc.add_table(rows=1, cols=len(headers))
    tbl.style = style
    hdr = tbl.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h
    for row in rows:
        cells = tbl.add_row().cells
        for i, v in enumerate(row):
            cells[i].text = "" if v is None else str(v)


def add_picture(doc: Document, path: str | Path, width_inches: float = 6.0) -> None:
    doc.add_picture(str(path), width=Inches(width_inches))


def save_doc(doc: Document, path: str | Path) -> None:
    Path(path).parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(path))


SEVEN_SECTIONS = [
    "Executive Summary",
    "Methodology",
    "Headline Findings",
    "Per-Cohort Observations",
    "Confounding Analysis",
    "Limitations",
    "Future Hypotheses",
]
