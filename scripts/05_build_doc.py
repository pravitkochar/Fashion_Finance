"""Phase 6 + 7 — build the Excel and Word deliverables.

Phase 6: reports/Fashion_Show_Event_Study.xlsx
  - Overview tab: 8 charts + summary tables
  - 15 per-company tabs: 4 charts + event-level table

Phase 7: reports/Findings.docx
  - 7 sections, >=4 embedded charts
"""
from __future__ import annotations

import json
import logging
import sys
from pathlib import Path

import matplotlib

matplotlib.use("Agg")
import matplotlib.pyplot as plt
import numpy as np
import pandas as pd
import seaborn as sns
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from openpyxl import Workbook
from openpyxl.drawing.image import Image as XLImage
from openpyxl.formatting.rule import ColorScaleRule
from openpyxl.styles import Alignment, Font, PatternFill
from openpyxl.utils import get_column_letter

logging.basicConfig(level=logging.INFO, format="%(asctime)s %(levelname)s %(message)s")
log = logging.getLogger("phase6_7")

ROOT = Path(__file__).resolve().parents[1]
DATA = ROOT / "data"
REPORTS = ROOT / "reports"
IMG = REPORTS / "img"

EVENTS = DATA / "events.csv"
CONFOUND = DATA / "confounders.csv"
PRICES = DATA / "prices_raw.csv"
SHOWS = DATA / "show_dates.csv"
FINDINGS = DATA / "phase1_findings.json"

XLSX = REPORTS / "Fashion_Show_Event_Study.xlsx"
DOCX = REPORTS / "Findings.docx"

TICKER_BENCH = {
    "MC.PA": "^FCHI", "KER.PA": "^FCHI", "RMS.PA": "^FCHI",
    "CPRI": "^GSPC", "TPR": "^GSPC", "COH": "^GSPC", "PVH": "^GSPC", "RL": "^GSPC",
    "BRBY.L": "^FTSE",
    "1913.HK": "^HSI",
    "BC.MI": "FTSEMIB.MI", "MONC.MI": "FTSEMIB.MI",
    "SFER.MI": "FTSEMIB.MI", "TOD.MI": "FTSEMIB.MI",
    "BOSS.DE": "^GDAXI",
}

ENTITY_TYPE = {
    "MC.PA": "Conglomerate", "KER.PA": "Conglomerate", "CPRI": "Conglomerate",
    "TPR": "Conglomerate", "COH": "Conglomerate", "PVH": "Conglomerate",
    "RMS.PA": "Pure-play", "BRBY.L": "Pure-play", "1913.HK": "Pure-play",
    "BC.MI": "Pure-play", "MONC.MI": "Pure-play", "BOSS.DE": "Pure-play",
    "RL": "Pure-play", "SFER.MI": "Pure-play", "TOD.MI": "Pure-play",
}

COMPANY_NAME = {
    "MC.PA": "LVMH", "KER.PA": "Kering", "CPRI": "Capri Holdings",
    "TPR": "Tapestry", "COH": "Coach Inc (legacy)", "PVH": "PVH Corp",
    "RMS.PA": "Hermes", "BRBY.L": "Burberry", "1913.HK": "Prada",
    "BC.MI": "Brunello Cucinelli", "MONC.MI": "Moncler", "BOSS.DE": "Hugo Boss",
    "RL": "Ralph Lauren", "SFER.MI": "Salvatore Ferragamo", "TOD.MI": "Tod's",
}

PRE = 30
POST = 30


def event_window(events: pd.DataFrame, prices: pd.DataFrame) -> dict:
    """Return dict event_id -> DataFrame with cols [t, AR] for events that fit."""
    by_ticker = {t: g.sort_values("date").reset_index(drop=True) for t, g in prices.groupby("ticker")}
    out: dict[str, pd.DataFrame] = {}
    for _, ev in events.iterrows():
        tk = ev["ticker"]
        bench = TICKER_BENCH.get(tk)
        if tk not in by_ticker or bench not in by_ticker:
            continue
        sdf = by_ticker[tk]
        bdf = by_ticker[bench][["date", "daily_return"]].rename(columns={"daily_return": "bench_ret"})
        td = sdf["date"].values
        sd = np.datetime64(pd.to_datetime(ev["show_date"]))
        idx = np.where(td >= sd)[0]
        if len(idx) == 0:
            continue
        i0 = int(idx[0])
        if i0 < PRE or i0 + POST >= len(sdf):
            continue
        win = sdf.iloc[i0 - PRE : i0 + POST + 1].copy()
        win["t"] = np.arange(-PRE, POST + 1)
        win = win.merge(bdf, on="date", how="left").dropna(subset=["daily_return", "bench_ret"])
        win["AR"] = win["daily_return"] - win["bench_ret"]
        out[ev["event_id"]] = win[["t", "AR", "date"]].reset_index(drop=True)
    return out


def make_overview_charts(df: pd.DataFrame, win_map: dict) -> list[Path]:
    IMG.mkdir(parents=True, exist_ok=True)
    paths = []

    ts = list(range(-PRE, POST + 1))
    rows = []
    for eid, w in win_map.items():
        for _, r in w.iterrows():
            rows.append({"event_id": eid, "t": int(r["t"]), "AR": float(r["AR"])})
    long = pd.DataFrame(rows)

    long_sorted = long.sort_values(["event_id", "t"])
    long_sorted["cum"] = long_sorted.groupby("event_id")["AR"].cumsum()
    pivot = (long_sorted
             .pivot(index="event_id", columns="t", values="cum")
             .reindex(columns=ts)
             .ffill(axis=1))
    mean_curve = pivot.mean(axis=0)
    p25 = pivot.quantile(0.25, axis=0)
    p75 = pivot.quantile(0.75, axis=0)

    p1 = IMG / "ov1_agg_car.png"
    plt.figure(figsize=(10, 5))
    plt.fill_between(ts, p25.values, p75.values, alpha=0.25, color="#3a6", label="25-75 pct")
    plt.plot(ts, mean_curve.values, color="#222", linewidth=1.6, label="Mean CAR")
    plt.axvline(0, color="grey", linewidth=0.7, linestyle="--")
    plt.axhline(0, color="grey", linewidth=0.7)
    plt.xlabel("Trading day"); plt.ylabel("CAR"); plt.title("1) Aggregate CAR (-30..+30) with 25-75 band")
    plt.legend(); plt.grid(alpha=0.3); plt.tight_layout(); plt.savefig(p1, dpi=140); plt.close()
    paths.append(p1)

    p2 = IMG / "ov2_heatmap.png"
    sub = long[(long["t"] >= -10) & (long["t"] <= 10)]
    eid_to_tk = dict(zip(df["event_id"], df["ticker"]))
    sub = sub.assign(ticker=sub["event_id"].map(eid_to_tk))
    heat = sub.groupby(["ticker", "t"])["AR"].mean().unstack().reindex(columns=range(-10, 11))
    plt.figure(figsize=(11, 6))
    sns.heatmap(heat, cmap="RdYlGn", center=0, cbar_kws={"label": "mean AR"})
    plt.title("2) Mean AR heatmap (rows=ticker, cols=trading day)")
    plt.xlabel("Trading day"); plt.ylabel("Ticker"); plt.tight_layout(); plt.savefig(p2, dpi=140); plt.close()
    paths.append(p2)

    p3 = IMG / "ov3_ar_t5_hist.png"
    x = df["AR_t5"].dropna() * 100
    plt.figure(figsize=(9, 5))
    plt.hist(x, bins=np.arange(np.floor(x.min()), np.ceil(x.max()) + 1, 1), edgecolor="white")
    plt.axvline(x.mean(), color="red", label=f"mean {x.mean():.2f}%")
    plt.axvline(x.median(), color="orange", linestyle="--", label=f"median {x.median():.2f}%")
    from scipy import stats as _st
    t = x.mean() / (x.std(ddof=1) / np.sqrt(len(x))) if x.std(ddof=1) else float("nan")
    plt.title(f"3) Single-event AR_t5 histogram (n={len(x)}, t={t:.2f})")
    plt.xlabel("AR_t5 (%)"); plt.ylabel("count"); plt.legend(); plt.grid(alpha=0.3)
    plt.tight_layout(); plt.savefig(p3, dpi=140); plt.close()
    paths.append(p3)

    p4 = IMG / "ov4_yearly.png"
    yr = df.groupby("year")["CAR_0to5"].mean()
    plt.figure(figsize=(10, 4.5))
    plt.plot(yr.index, yr.values * 100, marker="o", color="#444")
    plt.axhline(0, color="grey", linewidth=0.7)
    plt.title("4) Yearly mean CAR_0to5"); plt.xlabel("Year"); plt.ylabel("Mean CAR_0to5 (%)")
    plt.grid(alpha=0.3); plt.tight_layout(); plt.savefig(p4, dpi=140); plt.close()
    paths.append(p4)

    p5 = IMG / "ov5_ss_vs_fw.png"
    plt.figure(figsize=(10, 5))
    for season, color in (("spring", "#1b8a5a"), ("fall", "#9b3a3a")):
        eids = df.loc[df["season"] == season, "event_id"]
        sub = pivot.loc[pivot.index.isin(eids)]
        if len(sub):
            plt.plot(ts, sub.mean(axis=0).values, label=f"SS ({len(sub)})" if season == "spring" else f"FW ({len(sub)})",
                     color=color)
    plt.axvline(0, color="grey", linewidth=0.7, linestyle="--"); plt.axhline(0, color="grey", linewidth=0.7)
    plt.title("5) Mean CAR — SS vs FW"); plt.xlabel("Trading day"); plt.ylabel("CAR")
    plt.legend(); plt.grid(alpha=0.3); plt.tight_layout(); plt.savefig(p5, dpi=140); plt.close()
    paths.append(p5)

    p6 = IMG / "ov6_purepay_vs_conglom.png"
    df = df.assign(entity_type=df["ticker"].map(ENTITY_TYPE))
    plt.figure(figsize=(10, 5))
    for typ, color in (("Pure-play", "#1c4e80"), ("Conglomerate", "#a85d5d")):
        eids = df.loc[df["entity_type"] == typ, "event_id"]
        sub = pivot.loc[pivot.index.isin(eids)]
        if len(sub):
            plt.plot(ts, sub.mean(axis=0).values, label=f"{typ} ({len(sub)})", color=color)
    plt.axvline(0, color="grey", linewidth=0.7, linestyle="--"); plt.axhline(0, color="grey", linewidth=0.7)
    plt.title("6) Mean CAR — Pure-play vs Conglomerate"); plt.xlabel("Trading day"); plt.ylabel("CAR")
    plt.legend(); plt.grid(alpha=0.3); plt.tight_layout(); plt.savefig(p6, dpi=140); plt.close()
    paths.append(p6)

    p7 = IMG / "ov7_confound_vs_clean.png"
    plt.figure(figsize=(10, 5))
    for flag, color, label in ((True, "#a85d5d", "Confounded"), (False, "#1c4e80", "Clean")):
        eids = df.loc[df["confounded"] == flag, "event_id"]
        sub = pivot.loc[pivot.index.isin(eids)]
        if len(sub):
            plt.plot(ts, sub.mean(axis=0).values, label=f"{label} ({len(sub)})", color=color)
    plt.axvline(0, color="grey", linewidth=0.7, linestyle="--"); plt.axhline(0, color="grey", linewidth=0.7)
    plt.title("7) Mean CAR — Confounded vs Clean"); plt.xlabel("Trading day"); plt.ylabel("CAR")
    plt.legend(); plt.grid(alpha=0.3); plt.tight_layout(); plt.savefig(p7, dpi=140); plt.close()
    paths.append(p7)

    p8 = IMG / "ov8_best_worst_20.png"
    sorted_df = df.dropna(subset=["CAR_0to5"]).sort_values("CAR_0to5")
    bottom = sorted_df.head(20)
    top = sorted_df.tail(20)
    fig, axes = plt.subplots(1, 2, figsize=(14, 6))
    for ax, sub, title, color in ((axes[0], bottom, "Worst 20 (CAR_0to5)", "#a85d5d"),
                                   (axes[1], top, "Best 20 (CAR_0to5)", "#1b8a5a")):
        labels = [f'{r["ticker"]} {r["season"]}{r["year"]}' for _, r in sub.iterrows()]
        ax.barh(labels, sub["CAR_0to5"].values * 100, color=color)
        ax.set_xlabel("CAR_0to5 (%)"); ax.set_title(title); ax.grid(alpha=0.3)
        ax.invert_yaxis()
    plt.tight_layout(); plt.savefig(p8, dpi=140); plt.close()
    paths.append(p8)

    return paths


def make_company_charts(ticker: str, df_tk: pd.DataFrame, win_map: dict) -> list[Path]:
    paths = []
    safe = ticker.replace(".", "_").replace("/", "_")

    eids = list(df_tk["event_id"])
    rows = []
    for eid in eids:
        if eid not in win_map:
            continue
        for _, r in win_map[eid].iterrows():
            rows.append({"event_id": eid, "t": int(r["t"]), "AR": float(r["AR"])})
    if rows:
        long = pd.DataFrame(rows).sort_values(["event_id", "t"])
        long["cum"] = long.groupby("event_id")["AR"].cumsum()
        car_per = (long.pivot(index="event_id", columns="t", values="cum")
                       .reindex(columns=range(-PRE, POST + 1))
                       .ffill(axis=1))
    else:
        car_per = pd.DataFrame(columns=range(-PRE, POST + 1))

    p1 = IMG / f"co_{safe}_1_curve.png"
    plt.figure(figsize=(9, 4.5))
    if len(car_per):
        m = car_per.mean(axis=0)
        p25 = car_per.quantile(0.25, axis=0)
        p75 = car_per.quantile(0.75, axis=0)
        plt.fill_between(m.index, p25.values, p75.values, alpha=0.25, color="#3a6")
        plt.plot(m.index, m.values, color="#222", linewidth=1.6)
    plt.axvline(0, color="grey", linewidth=0.7, linestyle="--")
    plt.axhline(0, color="grey", linewidth=0.7)
    plt.title(f"{ticker} — Mean CAR (-30..+30) w/ 25-75 band"); plt.xlabel("Trading day"); plt.ylabel("CAR")
    plt.grid(alpha=0.3); plt.tight_layout(); plt.savefig(p1, dpi=140); plt.close()
    paths.append(p1)

    p2 = IMG / f"co_{safe}_2_hist.png"
    x = df_tk["CAR_0to5"].dropna() * 100
    plt.figure(figsize=(8, 4))
    if len(x):
        plt.hist(x, bins=12, edgecolor="white")
        plt.axvline(x.mean(), color="red", label=f"mean {x.mean():.2f}%")
        plt.axvline(x.median(), color="orange", linestyle="--", label=f"med {x.median():.2f}%")
        plt.legend()
    plt.title(f"{ticker} — CAR_0to5 distribution (n={len(x)})")
    plt.xlabel("CAR_0to5 (%)"); plt.grid(alpha=0.3); plt.tight_layout(); plt.savefig(p2, dpi=140); plt.close()
    paths.append(p2)

    p3 = IMG / f"co_{safe}_3_yearly.png"
    yr = df_tk.groupby("year").agg(n=("event_id", "count"), mean_car=("CAR_0to5", "mean"))
    fig, ax1 = plt.subplots(figsize=(9, 4.5))
    ax1.bar(yr.index, yr["n"].values, color="#bbb", label="# shows")
    ax1.set_ylabel("# shows"); ax1.set_xlabel("Year")
    ax2 = ax1.twinx()
    ax2.plot(yr.index, yr["mean_car"].values * 100, color="#a31", marker="o", label="mean CAR_0to5 (%)")
    ax2.axhline(0, color="grey", linewidth=0.6)
    ax2.set_ylabel("mean CAR_0to5 (%)")
    plt.title(f"{ticker} — Annual breakdown"); plt.tight_layout(); plt.savefig(p3, dpi=140); plt.close()
    paths.append(p3)

    p4 = IMG / f"co_{safe}_4_scatter.png"
    plt.figure(figsize=(9, 4))
    for season, color in (("spring", "#1b8a5a"), ("fall", "#1c4e80")):
        sub = df_tk[df_tk["season"] == season]
        if len(sub):
            plt.scatter(pd.to_datetime(sub["show_date"]), sub["AR_t5"] * 100, color=color, label=season, s=22)
    plt.axhline(0, color="grey", linewidth=0.7)
    plt.title(f"{ticker} — AR_t5 by event date"); plt.xlabel("Show date"); plt.ylabel("AR_t5 (%)")
    plt.legend(); plt.grid(alpha=0.3); plt.tight_layout(); plt.savefig(p4, dpi=140); plt.close()
    paths.append(p4)

    return paths


def autosize(ws, max_w: int = 28) -> None:
    for col in ws.columns:
        letter = get_column_letter(col[0].column)
        w = 8
        for c in col:
            v = c.value
            if v is None:
                continue
            w = max(w, min(max_w, len(str(v)) + 1))
        ws.column_dimensions[letter].width = w


def write_overview(wb: Workbook, df: pd.DataFrame, findings: dict, paths: list[Path]) -> None:
    ws = wb.active
    ws.title = "Overview"
    header_fill = PatternFill("solid", fgColor="0B2545")
    white_bold = Font(color="FFFFFF", bold=True)

    ws["A1"] = "Fashion Show Event Study — Phase 1 Overview"
    ws["A1"].font = Font(size=16, bold=True)
    ws.merge_cells("A1:H1")

    ws["A3"] = "Total events"; ws["B3"] = int(findings.get("n_events", len(df)))
    ws["A4"] = "Companies"; ws["B4"] = int(findings.get("n_companies", df["ticker"].nunique()))
    ws["A5"] = "Confounded events (earnings within 10d)"; ws["B5"] = int(df["confounded"].sum())
    ws["A6"] = "Confounder unknown (no earnings data)"; ws["B6"] = int(df["confounder_unknown"].sum())
    ws["A7"] = "Any test flagged"; ws["B7"] = "YES" if findings["any_flagged"] else "NO"

    ws["A9"] = "Test 1 — Aggregate t-stat"; ws["A9"].font = Font(bold=True, size=12)
    ws.append([""])
    headers = ["Window", "n", "mean CAR", "t-stat", "p-value", "flagged"]
    ws.append(headers)
    h_row = ws.max_row
    for c, h in enumerate(headers, start=1):
        cell = ws.cell(row=h_row, column=c)
        cell.fill = header_fill; cell.font = white_bold; cell.alignment = Alignment(horizontal="center")
    for r in findings["test1"]:
        ws.append([r["window"], r["n"], round(r["mean_CAR"], 4), round(r["t_stat"], 3),
                   round(r["p_value"], 4), "Y" if r["flagged"] else ""])

    ws.append([""])
    ws.append(["Test 2 — Median brand-level effects (|median CAR_0to5| > 3%)"])
    ws.cell(row=ws.max_row, column=1).font = Font(bold=True, size=12)
    ws.append(["Ticker", "Company", "n", "median CAR_0to5", "flagged"])
    h_row = ws.max_row
    for c, h in enumerate(["Ticker", "Company", "n", "median CAR_0to5", "flagged"], start=1):
        cell = ws.cell(row=h_row, column=c)
        cell.fill = header_fill; cell.font = white_bold; cell.alignment = Alignment(horizontal="center")
    for r in findings["test2"]["company_results"]:
        ws.append([r["ticker"], COMPANY_NAME.get(r["ticker"], r["ticker"]),
                   int(r["n_events"]), round(r["median_CAR_0to5"], 4),
                   "Y" if r["flagged"] else ""])

    ws.append([""])
    ws.append(["Test 3 — Peak deflection of aggregate CAR curve"])
    ws.cell(row=ws.max_row, column=1).font = Font(bold=True, size=12)
    ws.append(["Peak |deflection|", round(findings["test3"]["peak_deflection"], 4),
               "flagged" if findings["test3"]["flagged"] else "not flagged"])

    img_row = ws.max_row + 3
    for i, p in enumerate(paths):
        try:
            img = XLImage(str(p))
            img.width = 720; img.height = 360
            anchor = f"A{img_row + i * 22}"
            ws.add_image(img, anchor)
        except Exception as e:
            log.warning("could not embed %s: %s", p, e)

    autosize(ws)


def write_company(wb: Workbook, ticker: str, df_tk: pd.DataFrame, paths: list[Path],
                  prices: pd.DataFrame) -> None:
    sheet_name = ticker[:31]
    ws = wb.create_sheet(title=sheet_name)
    header_fill = PatternFill("solid", fgColor="0B2545")
    white_bold = Font(color="FFFFFF", bold=True)

    ws["A1"] = f"{COMPANY_NAME.get(ticker, ticker)} ({ticker})"
    ws["A1"].font = Font(size=14, bold=True); ws.merge_cells("A1:F1")
    ws["A2"] = f"Benchmark: {TICKER_BENCH.get(ticker, '?')}    Type: {ENTITY_TYPE.get(ticker, '?')}"
    listing = prices[prices["ticker"] == ticker]["date"].min()
    ws["A3"] = f"Listing/first price: {pd.to_datetime(listing).date() if pd.notna(listing) else 'n/a'}"
    ws["A4"] = f"Events: {len(df_tk)}"
    if len(df_tk):
        ws["A5"] = (f"median CAR_0to5: {df_tk['CAR_0to5'].median():.4f}    "
                    f"mean CAR_0to5: {df_tk['CAR_0to5'].mean():.4f}    "
                    f"confounded: {int(df_tk['confounded'].sum())}/{len(df_tk)}")

    img_row = 7
    for i, p in enumerate(paths):
        try:
            img = XLImage(str(p))
            img.width = 640; img.height = 320
            anchor = f"A{img_row + i * 18}"
            ws.add_image(img, anchor)
        except Exception as e:
            log.warning("could not embed %s: %s", p, e)

    table_row = img_row + len(paths) * 18 + 2
    cols = ["event_id", "brand_slug", "season", "year", "show_date", "designer",
            "trading_day_t0", "local_index_used",
            "CAR_pre30", "CAR_pre10", "CAR_pre5",
            "CAR_0to1", "CAR_0to5", "CAR_0to10", "CAR_0to30",
            "AR_t1", "AR_t5", "AR_t10",
            "earnings_within_10d", "confounder_unknown", "confounded"]
    for c, h in enumerate(cols, start=1):
        cell = ws.cell(row=table_row, column=c, value=h)
        cell.fill = header_fill; cell.font = white_bold; cell.alignment = Alignment(horizontal="center")
    df_out = df_tk.reindex(columns=cols).sort_values("show_date")
    for _, r in df_out.iterrows():
        table_row += 1
        for c, k in enumerate(cols, start=1):
            v = r[k]
            if isinstance(v, float):
                v = round(v, 4)
            ws.cell(row=table_row, column=c, value=v)

    n_rows = len(df_out)
    if n_rows:
        first_data_row = table_row - n_rows + 1
        last_data_row = table_row
        car_cols = [cols.index(c) + 1 for c in
                    ["CAR_pre30", "CAR_pre10", "CAR_pre5", "CAR_0to1", "CAR_0to5", "CAR_0to10", "CAR_0to30",
                     "AR_t1", "AR_t5", "AR_t10"]]
        for ci in car_cols:
            letter = get_column_letter(ci)
            rng = f"{letter}{first_data_row}:{letter}{last_data_row}"
            rule = ColorScaleRule(start_type="num", start_value=-0.05, start_color="C00000",
                                   mid_type="num", mid_value=0, mid_color="FFFFFF",
                                   end_type="num", end_value=0.05, end_color="00B050")
            ws.conditional_formatting.add(rng, rule)

    autosize(ws)


def build_doc(findings: dict, df: pd.DataFrame, ov_paths: list[Path]) -> None:
    doc = Document()

    h = doc.add_heading("Fashion Show Event Study — Phase 1 Findings", level=0)

    doc.add_heading("1. Executive Summary", level=1)
    flagged = findings["any_flagged"]
    summary = (
        f"Phase 1 analyzed {findings['n_events']} runway events across "
        f"{findings['n_companies']} listed fashion companies, 2000-2025. "
        + ("At least one of the three Phase 1 tests flagged a meaningful effect."
           if flagged else
           "None of the three Phase 1 tests flagged a statistically meaningful effect.")
        + " Test 1 (aggregate t-stat across 4 windows) "
        + ("identified " + ", ".join(t["window"] for t in findings["test1"] if t["flagged"]) +
           " as significant. " if any(t["flagged"] for t in findings["test1"]) else "showed no significant aggregate effect. ")
        + f"Test 2 (median brand-level effect, |·|>3%) flagged {findings['test2']['n_flagged']} of "
          f"{len(findings['test2']['company_results'])} companies. "
        + f"Test 3 (visual peak deflection) measured {findings['test3']['peak_deflection']*100:.2f}% peak "
        + ("(flagged, >1.5%)." if findings["test3"]["flagged"] else "(below 1.5% threshold).")
    )
    doc.add_paragraph(summary)

    doc.add_heading("2. Methodology", level=1)
    doc.add_paragraph(
        "Universe of 15 publicly-listed fashion houses across France (LVMH, Kering, Hermes), "
        "the US (Capri Holdings, Tapestry, Coach Inc legacy, PVH, Ralph Lauren), the UK (Burberry), "
        "Hong Kong (Prada), Italy (Brunello Cucinelli, Moncler, Salvatore Ferragamo, Tod's), and "
        "Germany (Hugo Boss). Events are SS and FW main women's ready-to-wear runway shows scraped "
        "from Vogue Runway 2000-2025. Couture, cruise/resort, pre-fall, and menswear shows are excluded."
    )
    doc.add_paragraph(
        "Event window is -30..+30 trading days around the first trading day on or after the show date. "
        "Abnormal returns are computed daily as stock log return minus local-index log return; CARs are "
        "the cumulative sum of abnormal returns over the specified subwindows. Earnings dates within "
        "10 trading days flag confounded events but do not exclude them."
    )

    doc.add_heading("3. Headline Findings", level=1)

    doc.add_heading("3.1 Test 1 — Aggregate t-stats", level=2)
    tbl = doc.add_table(rows=1, cols=6)
    tbl.style = "Light Grid Accent 1"
    hdr = tbl.rows[0].cells
    for i, h in enumerate(["Window", "n", "mean CAR", "t-stat", "p-value", "flagged"]):
        hdr[i].text = h
    for r in findings["test1"]:
        row = tbl.add_row().cells
        row[0].text = r["window"]; row[1].text = str(r["n"])
        row[2].text = f"{r['mean_CAR']:.4f}"; row[3].text = f"{r['t_stat']:.3f}"
        row[4].text = f"{r['p_value']:.4f}"; row[5].text = "Y" if r["flagged"] else ""

    doc.add_heading("3.2 Test 2 — Median brand-level effects", level=2)
    flagged_co = [r for r in findings["test2"]["company_results"] if r["flagged"]]
    if flagged_co:
        names = [f"{r['ticker']} ({r['median_CAR_0to5']*100:+.2f}%)" for r in flagged_co]
        doc.add_paragraph("Flagged companies (|median CAR_0to5| > 3%): " + ", ".join(names) + ".")
    else:
        doc.add_paragraph("No company exceeded the 3% median CAR_0to5 threshold.")

    doc.add_heading("3.3 Test 3 — Aggregate CAR curve", level=2)
    doc.add_paragraph(
        f"Peak absolute deflection of the mean CAR curve over the -30..+30 window: "
        f"{findings['test3']['peak_deflection']*100:.2f}% "
        f"({'flagged, > 1.5% threshold' if findings['test3']['flagged'] else 'below 1.5% threshold'})."
    )
    if ov_paths:
        doc.add_picture(str(ov_paths[0]), width=Inches(6.0))

    doc.add_heading("4. Per-Company Observations", level=1)
    co = pd.DataFrame(findings["test2"]["company_results"]).copy()
    co["abs_med"] = co["median_CAR_0to5"].abs()
    co = co.sort_values("abs_med", ascending=False)
    top = co.head(5)
    bullets = []
    for _, r in top.iterrows():
        nm = COMPANY_NAME.get(r["ticker"], r["ticker"])
        bullets.append(f"{nm} ({r['ticker']}): n={int(r['n_events'])}, median CAR_0to5={r['median_CAR_0to5']*100:+.2f}%")
    doc.add_paragraph("Strongest signals (largest |median CAR_0to5|):")
    for b in bullets:
        doc.add_paragraph(b, style="List Bullet")

    tbl = doc.add_table(rows=1, cols=4)
    tbl.style = "Light Grid Accent 1"
    h = tbl.rows[0].cells
    for i, x in enumerate(["Ticker", "Company", "n", "median CAR_0to5"]):
        h[i].text = x
    for _, r in co.sort_values("ticker").iterrows():
        row = tbl.add_row().cells
        row[0].text = r["ticker"]
        row[1].text = COMPANY_NAME.get(r["ticker"], r["ticker"])
        row[2].text = str(int(r["n_events"]))
        row[3].text = f"{r['median_CAR_0to5']*100:+.2f}%"

    doc.add_heading("5. Confounding Analysis", level=1)
    n_conf = int(df["confounded"].sum())
    n_clean = int((~df["confounded"]).sum())
    n_unk = int(df["confounder_unknown"].sum())
    doc.add_paragraph(
        f"Of {len(df)} events, {n_conf} fall within 10 trading days of an earnings release "
        f"(flagged confounded), {n_clean} are clean, and {n_unk} have unknown earnings status. "
        f"The chart below compares mean CAR curves for confounded vs clean subsets; if the runway "
        f"signal survives the earnings filter, the curves should track."
    )
    confound_chart = next((p for p in ov_paths if "confound" in p.name), None)
    if confound_chart:
        doc.add_picture(str(confound_chart), width=Inches(6.0))

    doc.add_heading("6. Limitations", level=1)
    for txt in [
        "Vogue Runway coverage is sparse for several brands pre-2003 and varies by brand throughout the period.",
        "Conglomerate tickers (LVMH, Kering, PVH, CPRI, TPR) capture the parent's full portfolio, not the runway brand alone — attribution dilution is unavoidable in Phase 1.",
        "Benchmark is the local broad-market index only. No sector or luxury-specific benchmark.",
        "Earnings dates from yfinance are best-effort; some tickers return no data and are flagged confounder_unknown.",
        "Phase 1 measures whether shows move stocks; it does not investigate the mechanism (sentiment, sales, leadership signal).",
        "Hugo Boss, Moncler, and Brunello Cucinelli have notably thin Vogue Runway coverage, lowering their event counts.",
    ]:
        doc.add_paragraph(txt, style="List Bullet")

    doc.add_heading("7. Phase 2 Hypotheses", level=1)
    for txt in [
        "Designer-change cohort: do CARs differ markedly around shows immediately following a creative-director appointment?",
        "Star-debut effect: isolate first-show-by-new-designer events vs steady-state shows.",
        "Social-media era cut: split sample at 2014 (Instagram-runway era) and compare.",
        "SS vs FW asymmetry: investigate whether one season's signal dominates and why (calendar, retail cycle).",
        "Pure-play vs conglomerate sensitivity: deeper test on dilution hypothesis.",
        "Cross-tabs by company size, market cap, and country exchange microstructure.",
    ]:
        doc.add_paragraph(txt, style="List Bullet")

    embedded = 1 if ov_paths else 0
    embedded += 1 if confound_chart else 0
    for p in ov_paths[:6]:
        if p == confound_chart:
            continue
        if embedded >= 4:
            break
        doc.add_picture(str(p), width=Inches(6.0))
        embedded += 1

    doc.save(DOCX)
    log.info("wrote %s", DOCX)


def main() -> int:
    REPORTS.mkdir(parents=True, exist_ok=True)
    IMG.mkdir(parents=True, exist_ok=True)

    events = pd.read_csv(EVENTS, parse_dates=["show_date", "trading_day_t0"])
    cf = pd.read_csv(CONFOUND)
    prices = pd.read_csv(PRICES, parse_dates=["date"])
    prices["date"] = prices["date"].dt.tz_localize(None).dt.normalize()
    findings = json.loads(FINDINGS.read_text())

    df = events.merge(cf, on=["event_id", "ticker"], how="left")
    df["confounded"] = df["confounded"].fillna(False)
    df["confounder_unknown"] = df["confounder_unknown"].fillna(False)

    log.info("building event windows for charts ...")
    win_map = event_window(events, prices)
    log.info("windows built: %d", len(win_map))

    log.info("rendering Overview charts ...")
    ov_paths = make_overview_charts(df, win_map)
    log.info("rendered %d overview charts", len(ov_paths))

    log.info("building Excel ...")
    wb = Workbook()
    write_overview(wb, df, findings, ov_paths)

    sorted_tickers = sorted(df["ticker"].unique(), key=lambda t: -len(df[df["ticker"] == t]))
    for tk in sorted_tickers:
        df_tk = df[df["ticker"] == tk].copy()
        co_paths = make_company_charts(tk, df_tk, win_map)
        write_company(wb, tk, df_tk, co_paths, prices)

    wb.save(XLSX)
    log.info("wrote %s", XLSX)

    log.info("building Word ...")
    build_doc(findings, df, ov_paths)
    log.info("Phase 6+7 complete.")
    return 0


if __name__ == "__main__":
    sys.exit(main())
